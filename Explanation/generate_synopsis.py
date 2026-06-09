from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# --- Title ---
title = doc.add_heading('FinRisk: AI-Powered Financial Risk Management System', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

# Subtitle
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('Project Synopsis')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x8A)

doc.add_paragraph()  # spacer

# --- 1. Project Overview ---
doc.add_heading('1. Project Overview', level=1)
doc.add_paragraph(
    'FinRisk is a state-of-the-art financial risk management dashboard that leverages machine learning '
    'and artificial intelligence to provide real-time risk assessment, portfolio intelligence, and '
    'actionable recommendations. The platform helps investors and institutions monitor seven distinct '
    'risk categories in a single, unified interface with a premium glassmorphic UI. It was developed '
    'as a team project for educational and research purposes in financial risk modeling.'
)

doc.add_paragraph(
    'Live Domain: https://finrisk.online'
)

# --- 2. Team ---
doc.add_heading('2. Project Team', level=1)
team_data = [
    ('Ayush Singh', 'Full-stack Development, ML Integration'),
    ('Aditya Singh', 'Backend Architecture, API Design'),
    ('Abhishek Kumar', 'Data Analysis, Model Training'),
    ('Bipin Singh', 'Frontend UI/UX, Visualization'),
]
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Team Member'
hdr[1].text = 'Role'
for name, role in team_data:
    row = table.add_row().cells
    row[0].text = name
    row[1].text = role

# --- 3. Technology Stack ---
doc.add_heading('3. Technology Stack', level=1)

doc.add_heading('3.1 Frontend', level=2)
frontend_data = [
    ('React 19.2.0', 'UI framework'),
    ('Vite 7.3.1', 'Build tool / dev server'),
    ('Tailwind CSS 3.4.13', 'Utility-first styling'),
    ('Recharts 3.8.0', 'Data visualization (charts, graphs)'),
    ('Framer Motion 12.38.0', 'Page transitions / animations'),
    ('Axios 1.13.6', 'HTTP client for API calls'),
    ('React Router DOM 7.13.1', 'Client-side routing'),
    ('Google OAuth', 'Social authentication'),
]
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Technology'
hdr[1].text = 'Purpose'
for tech, purpose in frontend_data:
    row = table.add_row().cells
    row[0].text = tech
    row[1].text = purpose

doc.add_heading('3.2 Backend', level=2)
backend_data = [
    ('FastAPI 0.135.1', 'Python REST framework'),
    ('SQLAlchemy 2.0.48', 'Database ORM'),
    ('PostgreSQL (Neon)', 'Serverless cloud database'),
    ('Uvicorn / Gunicorn', 'ASGI / WSGI servers'),
    ('Google OAuth / bcrypt', 'Authentication & security'),
    ('yfinance', 'Real-time market data'),
    ('OpenRouter API (GPT-4o-mini)', 'AI-powered analysis'),
    ('Pydantic', 'Data validation'),
]
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Technology'
hdr[1].text = 'Purpose'
for tech, purpose in backend_data:
    row = table.add_row().cells
    row[0].text = tech
    row[1].text = purpose

doc.add_heading('3.3 Machine Learning & Data Science', level=2)
ml_data = [
    ('XGBoost 3.2.0', 'Gradient boosting models'),
    ('CatBoost 1.2.7', 'Credit risk model (production)'),
    ('scikit-learn 1.8.0', 'Preprocessing, metrics'),
    ('pandas / numpy', 'Data manipulation & computing'),
    ('scipy', 'Scientific computing'),
    ('joblib', 'Model serialization'),
    ('imbalanced-learn', 'SMOTE / class weighting'),
]
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Library'
hdr[1].text = 'Purpose'
for lib, purpose in ml_data:
    row = table.add_row().cells
    row[0].text = lib
    row[1].text = purpose

# --- 4. Architecture ---
doc.add_heading('4. System Architecture', level=1)
doc.add_paragraph(
    'The application follows a modern three-tier architecture:'
)
arch_items = [
    ('Presentation Layer', 'React SPA hosted on Vercel with a glassmorphic UI, dark mode support, '
     'and responsive design. Communicates with the backend via Axios HTTP calls.'),
    ('Application Layer', 'FastAPI backend deployed on Render, exposing 25+ REST endpoints. '
     'Handles authentication, business logic, ML model inference, and AI-powered insights.'),
    ('Data Layer', 'PostgreSQL (Neon cloud) for persistent storage via SQLAlchemy ORM. '
     'Pickled ML model files (*.pkl) loaded at runtime for predictions. Yahoo Finance API '
     'for live market data.'),
]
for title, desc in arch_items:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}: ')
    run.bold = True
    p.add_run(desc)

doc.add_paragraph(
    'Data Flow: User Browser -> React Frontend (Vercel) -> FastAPI Backend (Render) -> '
    'PostgreSQL / ML Models / Yahoo Finance / OpenRouter API -> JSON Response'
)

# --- 5. Risk Modules ---
doc.add_heading('5. Risk Modules & ML Models', level=1)

risk_modules = [
    ('Credit Risk', 'CatBoost (prod), XGBoost', 'AUC: 0.771',
     'Predicts loan default probability using 120+ engineered features including income, '
     'credit amount, EXT_SOURCE scores, employment status, and demographics.'),
    ('Market Risk', 'ML VaR Estimator', 'RMSE: 0.012',
     'Estimates Value-at-Risk using 15 global market indices, commodities, and currencies '
     '(NIFTY, S&P 500, Gold, USD/INR, etc.). Integrated with live Yahoo Finance data.'),
    ('Business Risk', 'XGBoost Classifier', 'Accuracy: 96.3%',
     'Assesses business risk level based on revenue stability, expenses, competition level, '
     'and growth rate metrics.'),
    ('Operational Risk', 'Balanced Logistic Regression', 'Accuracy: 85.5%',
     'Scores operational risk using transaction type, amount, currency, counterparty, '
     'and category features.'),
    ('Liquidity Risk', 'Multi-class Classifier', 'Accuracy: 88.9%',
     'Classifies liquidity risk into 5 levels using 25 banking liquidity features '
     '(loans, deposits, T-bills, NPLs, interbank metrics).'),
    ('Financial Risk', 'Balanced Classifier', 'Accuracy: 97.3%',
     'Predicts bankruptcy probability using financial ratios (ROA, Leverage, Asset Turnover, '
     'Gross Profit/Liabilities, Operating Profit).'),
    ('E-Commerce Fraud', 'XGBoost Classifier', 'N/A',
     'Detects fraudulent transactions using amount, payment method, product category, '
     'location, device, and customer age features.'),
]

table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Module'
hdr[1].text = 'Model'
hdr[2].text = 'Performance'
hdr[3].text = 'Description'
for mod, model, perf, desc in risk_modules:
    row = table.add_row().cells
    row[0].text = mod
    row[1].text = model
    row[2].text = perf
    row[3].text = desc

# --- 6. API Endpoints ---
doc.add_heading('6. API Endpoints', level=1)
doc.add_paragraph(
    'The FastAPI backend exposes 25+ REST endpoints organized across the following groups:'
)
endpoints = [
    'Health: GET /',
    'Authentication: POST /auth/google, GET /profile, PUT /profile',
    'Portfolio: POST /portfolio/add, GET /portfolio/get/{email}, DELETE /portfolio/{id}/{email}',
    'Dashboard: GET /dashboard/stats?email=',
    'Credit Risk: POST /predict_credit_risk, GET /credit_predictions?email=',
    'Market Risk: POST /predict_market_risk, GET /market_features, GET /market_live_data, GET /market_risk_history',
    'Business Risk: POST /predict_business_risk, GET /business_features, GET /business_risk_history',
    'Liquidity Risk: POST /liquidity/predict, GET /liquidity/features, GET /liquidity/history',
    'Financial Risk: POST /financial/predict, GET /financial/history',
    'Fraud Detection: POST /predict_fraud, GET /fraud_history',
    'Market Data: GET /market-data',
    'AI Insights: GET /ai-insights, GET/POST /ai-risk-alerts',
    'Notifications: GET /notifications?email=',
]
for ep in endpoints:
    doc.add_paragraph(ep, style='List Bullet')

# --- 7. Project Structure ---
doc.add_heading('7. Project Directory Structure', level=1)
structure = [
    'Backend/ — FastAPI Python backend (main.py, database.py, models.py, risk API modules)',
    'React Frontend/ — React SPA (Vite, Tailwind CSS, Recharts, Framer Motion)',
    'Models/ — 21 trained ML artifacts (*.pkl files)',
    'Notebooks/ — 7 Jupyter notebooks for model training and research',
    'Datasets/ — Training data (Home Credit Default Risk dataset, 158 MB)',
    'Streamlit/ — Legacy prototype apps (6 Streamlit applications)',
    'Tests/ — pytest integration tests (20+ tests) and Locust load tests',
]
for item in structure:
    doc.add_paragraph(item, style='List Bullet')

# --- 8. Key Features ---
doc.add_heading('8. Key Features', level=1)
features = [
    'Seven distinct risk assessment modules in a unified dashboard',
    'Real-time market data integration via Yahoo Finance',
    'AI-powered portfolio insights using GPT-4o-mini via OpenRouter',
    'Google OAuth and email/password authentication',
    'Portfolio management with performance analytics and correlation matrices',
    'Dark mode with glassmorphic UI design',
    'Responsive design with animated page transitions',
    'Notification system for risk alerts',
    'Comprehensive test suite (integration + load testing)',
    'Production deployment on Vercel + Render with Neon PostgreSQL',
]
for f in features:
    doc.add_paragraph(f, style='List Bullet')

# --- 9. Deployment ---
doc.add_heading('9. Deployment', level=1)
doc.add_paragraph(
    'The project is deployed in production with the following infrastructure:'
)
deploy = [
    ('Frontend', 'Vercel — SPA with security headers and rewrites'),
    ('Backend', 'Render — FastAPI with Gunicorn/Uvicorn'),
    ('Database', 'Neon — Serverless PostgreSQL'),
    ('Live URL', 'https://finrisk.online'),
]
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Component'
hdr[1].text = 'Platform'
for comp, plat in deploy:
    row = table.add_row().cells
    row[0].text = comp
    row[1].text = plat

# --- 10. Conclusion ---
doc.add_heading('10. Conclusion', level=1)
doc.add_paragraph(
    'FinRisk demonstrates the practical application of machine learning and AI in financial risk '
    'management. By integrating seven distinct risk models into a single, production-grade web '
    'application with a modern UI, the project showcases end-to-end development capabilities '
    'spanning data science, backend engineering, frontend development, and cloud deployment. '
    'The system is fully functional, live at finrisk.online, and serves as a comprehensive '
    'platform for financial risk assessment and portfolio intelligence.'
)

# Save
output_path = '/Users/ayushsingh/Deep Learning/Financial Risk Management Project/Explanation/FinRisk_Synopsis.docx'
doc.save(output_path)
print(f'Saved to: {output_path}')
